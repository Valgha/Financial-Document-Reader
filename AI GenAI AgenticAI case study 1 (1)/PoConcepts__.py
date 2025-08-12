# import streamlit as st
# import os, re, base64
# import pdfplumber
# import docx
# from pathlib import Path
# from transformers import pipeline
# from PyPDF2 import PdfReader
# import re

# # Load NER pipeline (HuggingFace model)
# ner = pipeline("ner", model="dslim/bert-base-NER", aggregation_strategy="simple")

# # -------- File Handling -------- #
# def extract_text_from_pdf(file_path):
#     text = ""
#     with pdfplumber.open(file_path) as pdf:
#         for page in pdf.pages:
#             text += page.extract_text() or ""
#     return text

# def extract_text_from_docx(file_path):
#     doc = docx.Document(file_path)
#     return "\n".join([para.text for para in doc.paragraphs])

# def extract_text_from_txt(file_path):
#     return Path(file_path).read_text(encoding="utf-8")


# # rule based 
# def rule_based_financial_entities(text):
#     rules = {
#         "Counterparty": r"\bBANK\s+[A-Z]+\b",
#         "Notional": r"\b(?:USD|EUR|INR)\s?[\d,.]+\s?(?:mio|million|billion)?\b|\b\d+\s?(?:mio|million|billion)\b",
#         "ISIN": r"\b[A-Z]{2}\d{9}[A-Z0-9]\b",
#         "Underlying": r"(?:FLOAT\s\d{2}/\d{2}/\d{2}|Allianz\s+SE.*?\))",
#         "Maturity": r"\b\d+Y\s+[A-Z]+\b|\b\d{2}\s\w+\s\d{4}\b",
#         "Bid": r"\bestr\+\d+\s?bps\b",
#         "Offer": r"(?<=offer\s)[^\n]+",
#         "PaymentFrequency": r"\bQuarterly\b|\bQuaterly\b",
#         "Trade Date": r"\b\d{1,2}\s\w+\s\d{4}\b",
#         "Termination Date": r"\b\d{1,2}\s\w+\s\d{4}\b"
#     }
    
#     results = {}
#     for key, pattern in rules.items():
#         match = re.search(pattern, text, re.IGNORECASE)
#         results[key] = match.group(0) if match else None
    
#     return {"chat": results}


# # -------- NER Extraction -------- #
# def ner_extract(text):
#     entities = ner(text)
#     return entities

# # -------- Streamlit UI -------- #
# st.title("Financial Document Reader - AI Augmented")

# uploaded_file = st.file_uploader("Upload financial document", type=["pdf", "docx", "txt"])

# if uploaded_file:
#     file_path = f"temp_{uploaded_file.name}"
#     with open(file_path, "wb") as f:
#         f.write(uploaded_file.getbuffer())

#     # Detect type
#     if uploaded_file.name.endswith(".pdf"):
#         text = extract_text_from_pdf(file_path)
#     elif uploaded_file.name.endswith(".docx"):
#         text = extract_text_from_docx(file_path)
#     else:
#         text = extract_text_from_txt(file_path)

#     st.subheader("Raw Extracted Text")
#     st.write(text[:2000] + "..." if len(text) > 2000 else text)

#     mode = st.selectbox("Select Processing Mode", ["Rule-based", "NER Model", "LLM Pipeline"])

#     if mode == "Rule-based":
#         results = rule_based_financial_entities(text)
#         st.json(results)

#     elif mode == "NER Model":
#         entities = ner_extract(text)
#         st.json(entities)

#     elif mode == "LLM Pipeline":
#         st.write("🔍 LLM & RAG pipeline processing .")

#     os.remove(file_path)









# # app.py
# import os
# import json
# import pdfplumber
# import streamlit as st
# from pathlib import Path
# from sentence_transformers import SentenceTransformer
# import numpy as np
# import faiss
# import re
# import textwrap
# from typing import List, Dict

# # generation options
# USE_OPENAI = bool(os.getenv("OPENAI_API_KEY"))

# if USE_OPENAI:
#     import openai
#     openai.api_key = os.getenv("OPENAI_API_KEY")
# else:
#     # local generator fallback
#     from transformers import AutoTokenizer, AutoModelForSeq2SeqLM, pipeline
#     # NOTE: flan-t5-large is just an example. You can pick another model.
#     SUMMARIZER_MODEL = "google/flan-t5-large"
#     tokenizer = AutoTokenizer.from_pretrained(SUMMARIZER_MODEL)
#     model = AutoModelForSeq2SeqLM.from_pretrained(SUMMARIZER_MODEL)
#     local_generator = pipeline("text2text-generation", model=model, tokenizer=tokenizer, device=-1)  # CPU

# # Embedding model (sentence-transformers)
# EMBED_MODEL_NAME = "sentence-transformers/all-MiniLM-L6-v2"
# embed_model = SentenceTransformer(EMBED_MODEL_NAME)

# # ---------- Utilities ----------
# def extract_text_from_pdf(path: str) -> str:
#     text = []
#     with pdfplumber.open(path) as pdf:
#         for page in pdf.pages:
#             page_text = page.extract_text()
#             if page_text:
#                 text.append(page_text)
#     return "\n\n".join(text)

# def chunk_text(text: str, chunk_size: int = 800, overlap: int = 100) -> List[str]:
#     words = text.split()
#     chunks = []
#     start = 0
#     while start < len(words):
#         end = min(start + chunk_size, len(words))
#         chunk = " ".join(words[start:end])
#         chunks.append(chunk)
#         if end == len(words):
#             break
#         start = end - overlap
#     return chunks

# def build_faiss_index(chunks: List[str], embed_model: SentenceTransformer):
#     embeddings = embed_model.encode(chunks, show_progress_bar=True, convert_to_numpy=True)
#     dim = embeddings.shape[1]
#     index = faiss.IndexFlatL2(dim)
#     index.add(embeddings)
#     return index, embeddings

# def retrieve_topk(index, embeddings, query: str, chunks: List[str], embed_model: SentenceTransformer, k: int = 5):
#     q_emb = embed_model.encode([query], convert_to_numpy=True)
#     D, I = index.search(q_emb, k)
#     results = [{"chunk": chunks[i], "score": float(D[0][j])} for j, i in enumerate(I[0]) if i < len(chunks)]
#     return results

# # ---------- LLM prompt & extract ----------
# ENTITY_LIST = [
#     "Counterparty",
#     "Notional",
#     "ISIN",
#     "Underlying",
#     "Maturity",
#     "Bid",
#     "Offer",
#     "PaymentFrequency",
#     "Trade Date",
#     "Termination Date",
#     "Coupon",
#     "Barrier",
#     "Initial Valuation Date",
#     "Valuation Date",
#     "Upfront Payment"
# ]

# def build_prompt(entities: List[str], context: str) -> str:
#     """
#     Create a clear instruction prompt for the LLM to output JSON.
#     """
#     ent_lines = "\n".join([f"- {e}" for e in entities])
#     prompt = f"""
# You are a precise financial extraction assistant. Given the context passages from a financial document, extract the values for the requested entities.
# Return a JSON object with the entity names as keys and the extracted value as string or null if not found. Do not include extraneous commentary.

# Requested entities:
# {ent_lines}

# Context (use it to find values):
# """
#     # truncate context to reasonable length if too long
#     if len(context) > 4500:
#         context = context[:4500] + "\n\n[TRUNCATED]"
#     prompt += "\n\n" + context + "\n\nJSON:"
#     return prompt

# def call_llm(prompt: str, max_tokens: int = 512) -> str:
#     if USE_OPENAI:
#         # use gpt-3.5-turbo if available; adapt model name as needed
#         try:
#             resp = openai.ChatCompletion.create(
#                 model="gpt-3.5-turbo",
#                 messages=[{"role":"user","content": prompt}],
#                 max_tokens=max_tokens,
#                 temperature=0.0,
#             )
#             return resp["choices"][0]["message"]["content"].strip()
#         except Exception as e:
#             st.error(f"OpenAI call failed: {e}")
#             return ""
#     else:
#         # fallback to local generator
#         out = local_generator(prompt, max_length=max_tokens, do_sample=False)
#         return out[0]["generated_text"]

# def parse_json_from_text(text: str) -> Dict:
#     # Try to find a JSON block in text. Fallback: try to build minimal JSON via regex
#     try:
#         # attempt to find {...}
#         json_start = text.find("{")
#         json_end = text.rfind("}")
#         if json_start != -1 and json_end != -1 and json_end > json_start:
#             json_text = text[json_start:json_end+1]
#             return json.loads(json_text)
#     except Exception:
#         pass
#     # Fallback simple heuristics: parse "Entity: value" lines
#     result = {}
#     for line in text.splitlines():
#         line = line.strip()
#         if ":" in line:
#             k,v = line.split(":",1)
#             k = k.strip().strip('"').strip("'")
#             v = v.strip().strip('"').strip("'")
#             if k in ENTITY_LIST:
#                 result[k] = v
#     return result

# def extract_entities_rag(full_text: str, entities: List[str]) -> Dict:
#     chunks = chunk_text(full_text, chunk_size=600, overlap=120)
#     index, embeddings = build_faiss_index(chunks, embed_model)
#     # Build a multi-part context by retrieving per-entity or one retrieval
#     # We'll retrieve one combined context using an instruction that mentions all entities
#     query = "Extract the following entities: " + ", ".join(entities)
#     retrieved = retrieve_topk(index, embeddings, query, chunks, embed_model, k=6)
#     combined_context = "\n\n---\n\n".join([r["chunk"] for r in retrieved])
#     prompt = build_prompt(entities, combined_context)
#     raw = call_llm(prompt, max_tokens=512)
#     parsed = parse_json_from_text(raw)
#     # ensure all entities present
#     output = {e: parsed.get(e) or None for e in entities}
#     # Optionally add provenance
#     output["_retrieved_chunks_count"] = len(retrieved)
#     return output

# # ---------- Streamlit UI ----------
# st.set_page_config(layout="wide", page_title="Financial PDF LLM Extractor")
# st.title("Financial Document Reader — LLM (RAG) Extractor")

# uploaded_file = st.file_uploader("Upload a PDF to process (for LLM pipeline)", type=["pdf"])

# if uploaded_file:
#     # save temp
#     tmp_path = Path("tmp_uploaded.pdf")
#     tmp_path.write_bytes(uploaded_file.getbuffer())

#     with st.spinner("Extracting text from PDF..."):
#         text = extract_text_from_pdf(str(tmp_path))

#     st.subheader("Extracted raw text (first 3000 chars)")
#     st.code(text[:3000] + ("\n\n... (truncated)" if len(text) > 3000 else ""))

#     st.sidebar.header("Extraction options")
#     chosen = st.sidebar.multiselect("Entities to extract", ENTITY_LIST, default=ENTITY_LIST)
#     topk = st.sidebar.slider("Retriever results (top k chunks)", 1, 12, 6)

#     if st.button("Run LLM RAG extract"):
#         with st.spinner("Building index and querying LLM... this may take a little while"):
#             # Run extraction
#             chunks = chunk_text(text, chunk_size=600, overlap=120)
#             idx, embs = build_faiss_index(chunks, embed_model)
#             # retrieve using combined query
#             query = "Extract the following entities: " + ", ".join(chosen)
#             retrieved = retrieve_topk(idx, embs, query, chunks, embed_model, k=topk)
#             context = "\n\n---\n\n".join([r["chunk"] for r in retrieved])
#             prompt = build_prompt(chosen, context)
#             raw_answer = call_llm(prompt, max_tokens=800)
#             extracted = parse_json_from_text(raw_answer)
#             # normalize output
#             final = {e: extracted.get(e) or None for e in chosen}
#             final["_raw_llm_output"] = raw_answer
#             final["_retrieved_chunks_count"] = len(retrieved)

#         st.subheader("Extraction result JSON")
#         st.json(final)

#         st.subheader("LLM raw output")
#         st.code(raw_answer)

#     tmp_path.unlink(missing_ok=True)
# else:
#     st.info("Upload a PDF to run the LLM + RAG extractor. For DOCX/TXT handling use your other app logic.")






# # app.py - merged Financial Document Reader (rule-based, NER, RAG/LLM, green-highlight OCR)
# import os
# import io
# import json
# import re
# import tempfile
# from pathlib import Path
# from typing import List, Dict

# import streamlit as st
# import pdfplumber
# import docx
# from PIL import Image
# import numpy as np

# # optional libs; guard in case not installed
# try:
#     import cv2
# except Exception as e:
#     cv2 = None
# try:
#     import pytesseract
# except Exception as e:
#     pytesseract = None

# # transformers / HF / embeddings / faiss / openai
# from sentence_transformers import SentenceTransformer
# try:
#     import faiss
# except Exception:
#     faiss = None

# # NER pipeline
# from transformers import pipeline
# ner = pipeline("ner", model="dslim/bert-base-NER", aggregation_strategy="simple")

# # LLM config
# USE_OPENAI = bool(os.getenv("AIzaSyBGnhsxVtbSSxrTV6dfy2P0GrAc-XnpqII"))
# if USE_OPENAI:
#     import openai
#     openai.api_key = os.getenv("AIzaSyBGnhsxVtbSSxrTV6dfy2P0GrAc-XnpqII")
# else:
#     # local generator fallback (CPU)
#     from transformers import AutoTokenizer, AutoModelForSeq2SeqLM, pipeline as hf_pipeline
#     SUMMARIZER_MODEL = "google/flan-t5-base"
#     tok = AutoTokenizer.from_pretrained(SUMMARIZER_MODEL)
#     model = AutoModelForSeq2SeqLM.from_pretrained(SUMMARIZER_MODEL)
#     local_generator = hf_pipeline("text2text-generation", model=model, tokenizer=tok, device=-1)

# # embeddings
# EMBED_MODEL_NAME = "sentence-transformers/all-MiniLM-L6-v2"
# embed_model = SentenceTransformer(EMBED_MODEL_NAME)

# # --- RULES ---
# ENTITY_LIST = [
#     "Counterparty",
#     "Notional",
#     "ISIN",
#     "Underlying",
#     "Maturity",
#     "Bid",
#     "Offer",
#     "PaymentFrequency",
#     "Trade Date",
#     "Termination Date",
#     "Coupon",
#     "Barrier",
#     "Initial Valuation Date",
#     "Valuation Date",
#     "Upfront Payment"
# ]

# RULES = {
#     "Counterparty": r"\bBANK\s+[A-Z]+\b",
#     "Notional": r"\b(?:USD|EUR|INR)\s?[\d,.]+\s?(?:mio|million|billion)?\b|\b\d+\s?(?:mio|million|billion)\b",
#     "ISIN": r"\b[A-Z]{2}\d{9}[A-Z0-9]\b",
#     "Underlying": r"(?:FLOAT\s\d{2}/\d{2}/\d{2}|Allianz\s+SE.*?\)|[A-Z][a-zA-Z]+\sSE\b)",
#     "Maturity": r"\b\d+Y\s+[A-Z]+\b|\b\d{1,2}\s\w+\s\d{4}\b",
#     "Bid": r"\bestr\+\d+\s?bps\b",
#     "Offer": r"(?<=offer\s)[^\n]+",
#     "PaymentFrequency": r"\bQuarterly\b|\bQuaterly\b",
#     "Trade Date": r"\b\d{1,2}\s\w+\s\d{4}\b",
#     "Termination Date": r"\b\d{1,2}\s\w+\s\d{4}\b",
#     "Coupon": r"\b0%|\d+%|\bCoupon\s*\(C\)\b",
#     "Barrier": r"\bBarrier\s*\(?B\)?\b|\b\d{1,3}%\b",
#     "Initial Valuation Date": r"Initial Valuation Date[:\s]*\d{1,2}\s\w+\s\d{4}",
#     "Valuation Date": r"Valuation Date[:\s]*\d{1,2}\s\w+\s\d{4}",
#     "Upfront Payment": r"\*\*\*TBD\*\*\*%.*|Upfront Payment[:\s]*[^\n]+"
# }

# # ---------- Utilities ----------
# def extract_text_from_pdf(path: str) -> str:
#     text_blocks = []
#     with pdfplumber.open(path) as pdf:
#         for i, page in enumerate(pdf.pages):
#             page_text = page.extract_text()
#             if page_text:
#                 text_blocks.append(page_text)
#     return "\n\n".join(text_blocks)

# def extract_text_from_docx(file_path: str) -> str:
#     doc = docx.Document(file_path)
#     return "\n".join([p.text for p in doc.paragraphs])

# def extract_text_from_txt(file_path: str) -> str:
#     return Path(file_path).read_text(encoding="utf-8")

# # chunking for RAG
# def chunk_text(text: str, chunk_size: int = 800, overlap: int = 120) -> List[str]:
#     words = text.split()
#     chunks = []
#     start = 0
#     while start < len(words):
#         end = min(start + chunk_size, len(words))
#         chunks.append(" ".join(words[start:end]))
#         if end == len(words):
#             break
#         start = end - overlap
#     return chunks

# def build_faiss_index(chunks: List[str]):
#     if faiss is None:
#         raise RuntimeError("faiss not installed. Please install faiss-cpu or faiss-gpu.")
#     embs = embed_model.encode(chunks, show_progress_bar=False, convert_to_numpy=True)
#     dim = embs.shape[1]
#     index = faiss.IndexFlatL2(dim)
#     index.add(embs)
#     return index, embs

# def retrieve_topk(index, embeddings, query: str, chunks: List[str], k: int = 5):
#     q_emb = embed_model.encode([query], convert_to_numpy=True)
#     D, I = index.search(q_emb, k)
#     results = []
#     for rank, idx in enumerate(I[0]):
#         if idx < len(chunks):
#             results.append({"chunk": chunks[idx], "score": float(D[0][rank])})
#     return results

# def build_prompt(entities: List[str], context: str) -> str:
#     ent_lines = "\n".join([f"- {e}" for e in entities])
#     prompt = f"""You are a precise financial extraction assistant. Given the context passages from a financial document, extract the values for the requested entities.
# Return a JSON object with the entity names as keys and the extracted value as string or null if not found. Do not include extraneous commentary.

# Requested entities:
# {ent_lines}

# Context:
# {context}

# Return only JSON."""
#     # keep prompt length reasonable
#     return prompt

# def call_llm(prompt: str, max_tokens: int = 512) -> str:
#     if USE_OPENAI:
#         resp = openai.ChatCompletion.create(
#             model="gpt-3.5-turbo",
#             messages=[{"role": "user", "content": prompt}],
#             max_tokens=max_tokens,
#             temperature=0.0,
#         )
#         return resp["choices"][0]["message"]["content"].strip()
#     else:
#         out = local_generator(prompt, max_length=max_tokens, do_sample=False)
#         return out[0]["generated_text"]

# def parse_json_from_text(text: str) -> Dict:
#     # attempt to extract JSON block
#     try:
#         start = text.find("{")
#         end = text.rfind("}")
#         if start != -1 and end != -1 and end > start:
#             j = json.loads(text[start:end+1])
#             return j
#     except Exception:
#         pass
#     # fallback: key: value lines
#     res = {}
#     for line in text.splitlines():
#         if ":" in line:
#             k, v = line.split(":", 1)
#             k = k.strip().strip('"').strip("'")
#             v = v.strip().strip('"').strip("'")
#             if k in ENTITY_LIST:
#                 res[k] = v
#     return res

# def extract_entities_rag(full_text: str, entities: List[str], topk: int = 6):
#     chunks = chunk_text(full_text, chunk_size=600, overlap=120)
#     idx, embs = build_faiss_index(chunks)
#     query = "Extract: " + ", ".join(entities)
#     retrieved = retrieve_topk(idx, embs, query, chunks, k=topk)
#     combined_context = "\n\n---\n\n".join([r["chunk"] for r in retrieved])
#     prompt = build_prompt(entities, combined_context)
#     raw = call_llm(prompt, max_tokens=700)
#     parsed = parse_json_from_text(raw)
#     output = {e: (parsed.get(e) if isinstance(parsed, dict) else None) for e in entities}
#     output["_raw_llm_output"] = raw
#     output["_retrieved_chunks_count"] = len(retrieved)
#     return output

# # rule-based extraction
# def rule_based_financial_entities(text: str):
#     results = {}
#     for key, pattern in RULES.items():
#         m = re.search(pattern, text, re.IGNORECASE)
#         results[key] = m.group(0).strip() if m else None
#     return results

# # NER extraction (returns list)
# def ner_extract(text: str):
#     return ner(text)

# # --- Green-highlight OCR helper (OpenCV + pytesseract) ---
# def extract_green_regions_pil_image(pil_img: Image.Image, debug=False) -> List[Image.Image]:
#     """
#     Detect green-highlighted regions in a PIL image and return cropped PIL images for each region.
#     Requires cv2 and pytesseract for better results.
#     """
#     if cv2 is None:
#         raise RuntimeError("OpenCV (cv2) is required for green-region detection.")
#     if pytesseract is None:
#         raise RuntimeError("pytesseract is required for OCR.")

#     # convert to numpy array BGR for cv2
#     img = np.array(pil_img.convert("RGB"))
#     img_bgr = cv2.cvtColor(img, cv2.COLOR_RGB2BGR)
#     hsv = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2HSV)

#     # green color range - tune if your highlight hue differs
#     lower_green = np.array([40, 40, 40])
#     upper_green = np.array([90, 255, 255])

#     mask = cv2.inRange(hsv, lower_green, upper_green)
#     # dilate to fill holes
#     kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (9, 9))
#     mask = cv2.morphologyEx(mask, cv2.MORPH_CLOSE, kernel, iterations=2)
#     # find contours
#     contours, _ = cv2.findContours(mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
#     crops = []
#     h, w = mask.shape
#     for cnt in contours:
#         x, y, cw, ch = cv2.boundingRect(cnt)
#         # filter very small regions
#         if cw < 20 or ch < 8:
#             continue
#         # expand a bit
#         pad = 4
#         x0 = max(0, x - pad)
#         y0 = max(0, y - pad)
#         x1 = min(w, x + cw + pad)
#         y1 = min(h, y + ch + pad)
#         crop = pil_img.crop((x0, y0, x1, y1))
#         crops.append(crop)
#     # optionally sort crops top-to-bottom
#     crops = sorted(crops, key=lambda im: im.getbbox()[1] if im.getbbox() else 0)
#     return crops

# def ocr_image_pil(pil_img: Image.Image) -> str:
#     if pytesseract is None:
#         raise RuntimeError("pytesseract not available.")
#     text = pytesseract.image_to_string(pil_img, lang='eng')
#     return text

# # --- Streamlit UI and routing ---
# st.set_page_config(layout="wide", page_title="Financial Document Reader - Unified")
# st.title("Financial Document Reader — Unified (TXT / DOCX / PDF / Images)")

# uploaded_files = st.file_uploader("Upload files (pdf, docx, txt, png, jpg)", accept_multiple_files=True,
#                                   type=["pdf", "docx", "txt", "png", "jpg", "jpeg"])

# if not uploaded_files:
#     st.info("Upload one or more files. TXT -> rule-based, DOCX -> NER, PDF -> LLM/RAG (with OCR fallback). Images -> green-highlight OCR (default).")
#     st.stop()

# # sidebar options
# st.sidebar.header("Options")
# green_only = st.sidebar.checkbox("Extract only green-highlighted text from images / scanned PDFs", value=True)
# pdf_topk = st.sidebar.slider("PDF retriever top-k", 1, 12, 6)
# entities_to_extract = st.sidebar.multiselect("Entities for LLM extraction (PDF)", ENTITY_LIST, default=ENTITY_LIST)

# results_all = {}

# for uploaded in uploaded_files:
#     fname = uploaded.name
#     ext = fname.split(".")[-1].lower()
#     st.header(f"File: {fname} ({ext})")

#     # save temp file
#     tmp = tempfile.NamedTemporaryFile(delete=False, suffix="."+ext)
#     tmp.write(uploaded.getbuffer())
#     tmp.flush()
#     tmp.close()
#     tmp_path = tmp.name

#     try:
#         if ext == "txt":
#             text = extract_text_from_txt(tmp_path)
#             st.subheader("Raw text (full)")
#             st.text(text)
#             st.subheader("Rule-based extraction")
#             rb = rule_based_financial_entities(text)
#             st.json(rb)
#             results_all[fname] = {"mode": "rule-based-txt", "results": rb}

#         elif ext == "docx":
#             text = extract_text_from_docx(tmp_path)
#             st.subheader("Extracted text (docx)")
#             st.text(text[:5000] + ("\n\n... (truncated)" if len(text) > 5000 else ""))
#             st.subheader("NER extraction (list of entities)")
#             ner_out = ner_extract(text)
#             st.json(ner_out)
#             results_all[fname] = {"mode": "ner-docx", "results": ner_out}

#         elif ext == "pdf":
#             # Extract text; if empty or user chooses green-only we use OCR / green-extract
#             extracted_text = extract_text_from_pdf(tmp_path)
#             if (not extracted_text.strip()) or green_only:
#                 # try OCR per page and/or green-highlight on page images
#                 try:
#                     ocr_full = []
#                     green_texts = []
#                     with pdfplumber.open(tmp_path) as pdf:
#                         for pnum, page in enumerate(pdf.pages):
#                             # get page image at high resolution
#                             try:
#                                 pil_page = page.to_image(resolution=150).original
#                             except Exception:
#                                 # fallback: render raster using PIL from PDF page crop (pdfplumber may not support)
#                                 pil_page = None
#                             if pil_page is None:
#                                 # fallback: attempt to extract text only
#                                 page_text = page.extract_text() or ""
#                                 if page_text.strip():
#                                     ocr_full.append(page_text)
#                                 continue

#                             if green_only:
#                                 if cv2 is None or pytesseract is None:
#                                     st.warning("OpenCV and pytesseract needed for green-only extraction; falling back to full OCR.")
#                                     page_text = pytesseract.image_to_string(pil_page) if pytesseract else ""
#                                     ocr_full.append(page_text)
#                                 else:
#                                     crops = extract_green_regions_pil_image(pil_page)
#                                     for c in crops:
#                                         txt = ocr_image_pil(c)
#                                         green_texts.append(txt)
#                             else:
#                                 # full OCR of page
#                                 page_text = ocr_image_pil(pil_page) if pytesseract else page.extract_text() or ""
#                                 ocr_full.append(page_text)

#                     combined = ""
#                     if green_only:
#                         combined = "\n".join(green_texts)
#                         st.subheader("Green-highlight OCR results (concatenated)")
#                         st.text(combined or "No green-highlight text found.")
#                         # run rule-based + ner heuristics on green text
#                         rb = rule_based_financial_entities(combined)
#                         st.subheader("Rule-based extraction from green text")
#                         st.json(rb)
#                         results_all[fname] = {"mode": "green-ocr-pdf", "results": rb, "raw_green_text": combined}
#                     else:
#                         # full OCR results
#                         combined = "\n".join(ocr_full)
#                         st.subheader("OCR extracted text (first 5000 chars)")
#                         st.text(combined[:5000] + ("\n\n... (truncated)" if len(combined) > 5000 else ""))
#                         # Use RAG on combined text
#                         st.subheader("Running LLM RAG extraction on OCR text")
#                         rag_out = extract_entities_rag(combined, entities_to_extract, topk=pdf_topk)
#                         st.json(rag_out)
#                         results_all[fname] = {"mode": "rag-pdf-ocr", "results": rag_out}
#                 except Exception as e:
#                     st.error(f"Error during OCR/green extraction: {e}")
#                     # fallback to text extraction if available
#                     if extracted_text.strip():
#                         st.subheader("Fallback: extracted PDF selectable text (first 3000 chars)")
#                         st.text(extracted_text[:3000] + ("\n\n... (truncated)" if len(extracted_text) > 3000 else ""))
#                         st.subheader("RAG extraction")
#                         rag_out = extract_entities_rag(extracted_text, entities_to_extract, topk=pdf_topk)
#                         st.json(rag_out)
#                         results_all[fname] = {"mode": "rag-pdf-fallback", "results": rag_out}
#             else:
#                 # we have selectable text and green_only not requested -> do RAG
#                 st.subheader("Extracted text (first 3000 chars)")
#                 st.text(extracted_text[:3000] + ("\n\n... (truncated)" if len(extracted_text) > 3000 else ""))
#                 st.subheader("Running LLM RAG extraction on selectable text")
#                 rag_out = extract_entities_rag(extracted_text, entities_to_extract, topk=pdf_topk)
#                 st.json(rag_out)
#                 results_all[fname] = {"mode": "rag-pdf-selectable", "results": rag_out}

#         elif ext in ("png", "jpg", "jpeg"):
#             pil = Image.open(tmp_path).convert("RGB")
#             st.image(pil, caption="Uploaded image", use_column_width=True)
#             if green_only:
#                 if cv2 is None or pytesseract is None:
#                     st.error("Green-only extraction requires OpenCV and pytesseract. Please install them.")
#                     # fallback to full OCR
#                     if pytesseract:
#                         full = ocr_image_pil(pil)
#                         st.subheader("Full OCR result")
#                         st.text(full)
#                         rb = rule_based_financial_entities(full)
#                         st.subheader("Rule-based extraction from full OCR")
#                         st.json(rb)
#                         results_all[fname] = {"mode": "ocr-image-fallback", "results": rb}
#                     else:
#                         results_all[fname] = {"mode": "image-no-ocr", "results": None}
#                 else:
#                     crops = extract_green_regions_pil_image(pil)
#                     green_texts = [ocr_image_pil(c) for c in crops]
#                     combined = "\n".join(green_texts)
#                     st.subheader("Green-highlight OCR (concatenated)")
#                     st.text(combined or "No green-highlight found.")
#                     rb = rule_based_financial_entities(combined)
#                     st.subheader("Rule-based extraction from green text")
#                     st.json(rb)
#                     results_all[fname] = {"mode": "green-ocr-image", "results": rb, "raw_green_text": combined}
#             else:
#                 # full OCR
#                 if pytesseract:
#                     full = ocr_image_pil(pil)
#                     st.subheader("Full OCR result")
#                     st.text(full[:5000] + ("\n\n... (truncated)" if len(full) > 5000 else ""))
#                     st.subheader("Rule-based extraction from full OCR")
#                     rb = rule_based_financial_entities(full)
#                     st.json(rb)
#                     results_all[fname] = {"mode": "ocr-image-full", "results": rb}
#                 else:
#                     st.error("pytesseract not available for image OCR.")
#                     results_all[fname] = {"mode": "image-no-ocr", "results": None}
#         else:
#             st.warning(f"Unsupported file type: {ext}")
#             results_all[fname] = {"mode": "unsupported", "results": None}
#     finally:
#         try:
#             os.unlink(tmp_path)
#         except Exception:
#             pass

# # Show summary
# st.sidebar.header("Processed files summary")
# st.sidebar.write({k: v["mode"] for k, v in results_all.items()})

# st.header("All results (JSON)")
# st.json(results_all)

# instead of green word ocr we want ocr of full pdf to give output as shown in the architecture



# import re
# import json
# from docx import Document
# from pdfminer.high_level import extract_text

# # Predefined target entities (exact schema from your green-highlighted example)
# TARGET_FIELDS = [
#     "Party A",
#     "Party B",
#     "Trade Date",
#     "Initial Valuation Date",
#     "Final Valuation Date",
#     "Settlement Date",
#     "Notional Amount",
#     "Currency",
#     "Coupon",
#     "Barrier",
#     "Strike Price",
#     "Underlying Asset",
#     "Exchange",
#     "Calculation Agent",
#     "Business Day Convention",
#     "Day Count Convention"
# ]

# # Regex patterns for each field (can be expanded based on document format)
# PATTERNS = {
#     "Party A": r"Party\s*A[:\s]*([A-Za-z0-9\s&.,-]+)",
#     "Party B": r"Party\s*B[:\s]*([A-Za-z0-9\s&.,-]+)",
#     "Trade Date": r"Trade Date[:\s]*([\w\s,0-9-]+)",
#     "Initial Valuation Date": r"Initial Valuation Date[:\s]*([\w\s,0-9-]+)",
#     "Final Valuation Date": r"Final Valuation Date[:\s]*([\w\s,0-9-]+)",
#     "Settlement Date": r"Settlement Date[:\s]*([\w\s,0-9-]+)",
#     "Notional Amount": r"Notional Amount[:\s]*([\w\s,.]+)",
#     "Currency": r"Currency[:\s]*([A-Za-z]+)",
#     "Coupon": r"Coupon[:\s]*([\w\s%.,-]+)",
#     "Barrier": r"Barrier[:\s]*([\w\s%.,-]+)",
#     "Strike Price": r"Strike Price[:\s]*([\w\s%.,-]+)",
#     "Underlying Asset": r"Underlying Asset[:\s]*([\w\s&.,-]+)",
#     "Exchange": r"Exchange[:\s]*([\w\s&.,-]+)",
#     "Calculation Agent": r"Calculation Agent[:\s]*([\w\s&.,-]+)",
#     "Business Day Convention": r"Business Day Convention[:\s]*([\w\s&.,-]+)",
#     "Day Count Convention": r"Day Count Convention[:\s]*([\w\s&.,-]+)"
# }

# def extract_text_from_file(file_path):
#     if file_path.lower().endswith(".docx"):
#         doc = Document(file_path)
#         return "\n".join([p.text for p in doc.paragraphs])
#     elif file_path.lower().endswith(".pdf"):
#         return extract_text(file_path)
#     else:
#         raise ValueError("Unsupported file type. Use DOCX or PDF.")

# def extract_entities(text):
#     extracted = {field: None for field in TARGET_FIELDS}
#     for field, pattern in PATTERNS.items():
#         match = re.search(pattern, text, re.IGNORECASE)
#         if match:
#             extracted[field] = match.group(1).strip()
#     return extracted

# if __name__ == "__main__":
#     file_path = "contract.docx"  # Change to your file
#     text_data = extract_text_from_file(file_path)
#     result = extract_entities(text_data)
    
#     # Save as JSON
#     with open("extracted_entities.json", "w") as f:
#         json.dump(result, f, indent=4)

#     print(json.dumps(result, indent=4))







# You are a precise financial extraction assistant. Given the context from a financial document, extract ONLY the values for the requested entities below:
# prompt = """
# "Identify the name of the given company named'Party A'and give the value of 'Party A' to 'Counterparty'",
# "Identify ' Initial Valuation Date' in the file",
# "identify ' Notional'in the file and give the value ",
# "Likewise for --  'Valuation Date', ' Maturity', 'Underlying', 'Coupon', ' Barrier'"
# "Identify the key 'Calendar 'and give the data for it"


from dotenv import load_dotenv
# Load environment variables from .env file
load_dotenv()

import os
import io
import json
import re
import tempfile
from pathlib import Path
from typing import List, Dict

import streamlit as st
import pdfplumber
import docx
from PIL import Image
import numpy as np

# Optional libraries
try:
    import cv2
except ImportError:
    cv2 = None
try:
    import pytesseract
except ImportError:
    pytesseract = None
try:
    from sentence_transformers import SentenceTransformer
    import faiss
    from transformers import pipeline, AutoTokenizer, AutoModelForSeq2SeqLM
except ImportError:
    st.warning("Some AI libraries are not installed. RAG/NER features will be disabled.")
    SentenceTransformer = None
    faiss = None
    pipeline = None
    AutoTokenizer = None
    AutoModelForSeq2SeqLM = None

# --- AI Model Initialization ---
ner = pipeline("ner", model="dslim/bert-base-NER", aggregation_strategy="simple") if pipeline else None

USE_OPENAI = bool(os.getenv("OPENAI_API_KEY"))
if USE_OPENAI:
    import openai
    openai.api_key = os.getenv("OPENAI_API_KEY")
    local_generator = None
    st.sidebar.success("OpenAI API key detected. Using fast RAG/LLM for PDFs and DOCXs.")
else:
    st.sidebar.warning("OpenAI API key not found. Falling back to local models (may be slower).")
    SUMMARIZER_MODEL = "google/flan-t5-base"
    try:
        if AutoTokenizer and AutoModelForSeq2SeqLM and pipeline:
            tok = AutoTokenizer.from_pretrained(SUMMARIZER_MODEL)
            model = AutoModelForSeq2SeqLM.from_pretrained(SUMMARIZER_MODEL)
            local_generator = pipeline("text2text-generation", model=model, tokenizer=tok, device=-1)
        else:
            local_generator = None
    except Exception as e:
        st.error(f"Could not load local LLM. Error: {e}")
        local_generator = None

embed_model = SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2") if SentenceTransformer else None

ENTITY_LIST = [
    "Counterparty", "Notional", "ISIN", "Underlying", "Maturity", "Bid",
    "Offer", "PaymentFrequency", "Trade Date", "Termination Date", "Coupon",
    "Barrier", "Initial Valuation Date", "Valuation Date", "Upfront Payment"
]

RULES = {
    "Counterparty": r"\bBANK\s+[A-Z]+\b",
    "Notional": r"\b(?:USD|EUR|INR)\s?[\d,.]+\s?(?:mio|million|billion)?\b|\b\d+\s?(?:mio|million|billion)\b",
    "ISIN": r"\b[A-Z]{2}\d{9}[A-Z0-9]\b",
    "Underlying": r"(?:FLOAT\s\d{2}/\d{2}/\d{2}|Allianz\s+SE.*?\)|[A-Z][a-zA-Z]+\sSE\b)",
    "Maturity": r"\b\d+Y\s+[A-Z]+\b|\b\d{1,2}\s\w+\s\d{4}\b",
    "Bid": r"\bestr\+\d+\s?bps\b",
    "Offer": r"(?<=offer\s)[^\n]+",
    "PaymentFrequency": r"\bQuarterly\b|\bQuaterly\b",
    "Trade Date": r"\b\d{1,2}\s\w+\s\d{4}\b",
    "Termination Date": r"\b\d{1,2}\s\w+\s\d{4}\b",
    "Coupon": r"\b0%|\d+%|\bCoupon\s*\(C\)\b",
    "Barrier": r"\bBarrier\s*\(?B\)?\b|\b\d{1,3}%\b",
    "Initial Valuation Date": r"Initial Valuation Date[:\s]*\d{1,2}\s\w+\s\d{4}",
    "Valuation Date": r"Valuation Date[:\s]*\d{1,2}\s\w+\s\d{4}",
    "Upfront Payment": r"\*\*\*TBD\*\*\*%.*|Upfront Payment[:\s]*[^\n]+"
}

# --- Utility Functions ---
@st.cache_data
def extract_text_from_pdf(path: str) -> str:
    text_blocks = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_blocks.append(page_text)
    return "\n\n".join(text_blocks)

def extract_text_from_docx(file_path: str) -> str:
    doc = docx.Document(file_path)
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text.strip())
    return "\n".join(text_parts)

def extract_text_from_txt(file_path: str) -> str:
    return Path(file_path).read_text(encoding="utf-8")

def chunk_text(text: str, chunk_size: int = 800, overlap: int = 120) -> List[str]:
    words = text.split()
    chunks = []
    start = 0
    while start < len(words):
        end = min(start + chunk_size, len(words))
        chunks.append(" ".join(words[start:end]))
        if end == len(words):
            break
        start = end - overlap
    return chunks

@st.cache_resource
def build_faiss_index(chunks: List[str]):
    if not (faiss and embed_model):
        st.warning("FAISS or Sentence-Transformers not installed. RAG features are disabled.")
        return None, None
    embs = embed_model.encode(chunks, show_progress_bar=False, convert_to_numpy=True)
    dim = embs.shape[1]
    index = faiss.IndexFlatL2(dim)
    index.add(embs)
    return index, embs

def retrieve_topk(index, query: str, chunks: List[str], k: int = 5):
    if not index:
        return [{"chunk": "FAISS index not available.", "score": 0}]
    q_emb = embed_model.encode([query], convert_to_numpy=True)
    D, I = index.search(q_emb, k)
    results = []
    for rank, idx in enumerate(I[0]):
        if idx < len(chunks):
            results.append({"chunk": chunks[idx], "score": float(D[0][rank])})
    return results

def build_prompt(entities: List[str], context: str) -> str:
    ent_lines = ",\n".join([f'"{e}": ""' for e in entities])
    return f"""
You are a precise financial extraction assistant.
Given the context from a financial document, extract ONLY the values for the requested entities below.

 
"Identify the name of the given company named'Party A'and give the value of 'Party A' to 'Counterparty'",
"Identify ' Initial Valuation Date' in the file",
"identify ' Notional'in the file and give the value ",
"Likewise for --  'Valuation Date', ' Maturity', 'Underlying', 'Coupon', ' Barrier'"
"Identify the key 'Calendar 'and give the data for it"

Return the result strictly in JSON format, with keys exactly as in the entity list. 
Do not include extra text, explanations, or keys.

JSON format:
key : value pair
"""

def call_llm(prompt: str, max_tokens: int = 512) -> str:
    if USE_OPENAI:
        if not openai.api_key:
            return "API key for OpenAI not set."
        try:
            resp = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=max_tokens,
                temperature=0.0,
            )
            return resp["choices"][0]["message"]["content"].strip()
        except Exception as e:
            return f"Error calling OpenAI API: {e}"
    else:
        if local_generator:
            try:
                out = local_generator(prompt, max_length=max_tokens, do_sample=False)
                return out[0]["generated_text"]
            except Exception as e:
                return f"Error calling local LLM: {e}"
        else:
            return "Local LLM generator not available."

def parse_json_from_text(text: str) -> Dict:
    try:
        start = text.find("{")
        end = text.rfind("}")
        if start != -1 and end != -1 and end > start:
            return json.loads(text[start:end+1])
    except Exception:
        pass
    return {}

def extract_entities_rag(full_text: str, entities: List[str], topk: int = 6):
    if not (faiss and embed_model and (USE_OPENAI or local_generator)):
        return {"error": "RAG dependencies or LLM not configured."}
    chunks = chunk_text(full_text, chunk_size=600, overlap=120)
    idx, _ = build_faiss_index(chunks)
    if not idx:
        return {"error": "Failed to build FAISS index."}
    query = "Extract: " + ", ".join(entities)
    retrieved = retrieve_topk(idx, query, chunks, k=topk)
    combined_context = "\n\n---\n\n".join([r["chunk"] for r in retrieved])
    prompt = build_prompt(entities, combined_context)
    raw = call_llm(prompt, max_tokens=700)
    parsed = parse_json_from_text(raw)
    return {"raw_output": raw, "parsed": {e: parsed.get(e) for e in entities}}

def remove_null_values(data: dict) -> dict:
    return {k: v for k, v in data.items() if v not in [None, "", "null", "Null", "NULL"]}

def rule_based_financial_entities(text: str):
    return {key: re.search(pattern, text, re.IGNORECASE).group(0).strip() if re.search(pattern, text, re.IGNORECASE) else None for key, pattern in RULES.items()}

def ner_extract(text: str):
    if not ner:
        return {"error": "NER pipeline not available."}
    return ner(text)

def ocr_image_pil(pil_img: Image.Image) -> str:
    if not pytesseract:
        raise RuntimeError("pytesseract is not available.")
    return pytesseract.image_to_string(pil_img, lang='eng')

# --- Streamlit UI ---
st.set_page_config(layout="wide", page_title="Financial Document Reader - Unified")
st.title("Financial Document Reader ")
st.markdown("A unified tool for extracting financial entities from various document formats.")

uploaded_files = st.file_uploader(
    "Upload your documents (pdf, docx, txt, png, jpg)", 
    accept_multiple_files=True, 
    type=["pdf", "docx", "txt", "png", "jpg", "jpeg"]
)

if not uploaded_files:
    st.info("Upload one or more files to get started.")
    st.stop()

st.sidebar.header("Options ⚙️")
green_only = st.sidebar.checkbox("Extract only green-highlighted text from images/scanned PDFs", value=False)
pdf_topk = st.sidebar.slider("PDF RAG retriever top-k", 1, 12, 6)
entities_to_extract = st.sidebar.multiselect("Entities for RAG extraction", ENTITY_LIST, default=ENTITY_LIST)

for uploaded in uploaded_files:
    fname, ext = uploaded.name, uploaded.name.split(".")[-1].lower()
    st.header(f"Processing: {fname} ({ext})")
    st.markdown("---")

    with tempfile.NamedTemporaryFile(delete=False, suffix="." + ext) as tmp:
        tmp.write(uploaded.getbuffer())
        tmp_path = tmp.name

    try:
        if ext == "txt":
            text = extract_text_from_txt(tmp_path)
            st.subheader("Raw Text")
            st.text(text)
            extracted = rule_based_financial_entities(text)
            st.subheader("Rule-based Extraction")
            st.json(remove_null_values(extracted))
            st.subheader("NER Extraction")
            st.json(ner_extract(text))
        
        elif ext == "docx":
            with st.spinner("Extracting text from DOCX..."):
                text = extract_text_from_docx(tmp_path)
            st.subheader("Extracted Text")
            st.text(text[:3000] + ("..." if len(text) > 3000 else ""))

            if USE_OPENAI or local_generator:
                with st.spinner("Running RAG-based extraction with LLM..."):
                    rag_out = extract_entities_rag(text, entities_to_extract, topk=pdf_topk)
                st.subheader("Raw LLM Output")
                st.text(rag_out["raw_output"])
                st.subheader("Parsed JSON Output")
                st.json(remove_null_values(rag_out["parsed"]))
            st.subheader("NER Output")
            st.json(ner_extract(text))
                
        elif ext == "pdf":
            with st.spinner("Extracting text from PDF..."):
                extracted_text = extract_text_from_pdf(tmp_path)

            if not extracted_text.strip() and pytesseract:
                st.warning("No selectable text found. Running OCR...")
                with pdfplumber.open(tmp_path) as pdf:
                    ocr_texts = []
                    for page in pdf.pages:
                        pil_page = page.to_image(resolution=150).original
                        ocr_texts.append(ocr_image_pil(pil_page))
                extracted_text = "\n".join(ocr_texts)

            if extracted_text.strip():
                st.subheader("Extracted Text")
                st.text(extracted_text[:3000] + ("..." if len(extracted_text) > 3000 else ""))

                if USE_OPENAI or local_generator:
                    with st.spinner("Running RAG-based extraction with LLM..."):
                        rag_out = extract_entities_rag(extracted_text, entities_to_extract, topk=pdf_topk)
                    st.subheader("Raw LLM Output")
                    st.text(rag_out["raw_output"])
                    st.subheader("Parsed JSON Output")
                    st.json(remove_null_values(rag_out["parsed"]))
                st.subheader("NER Output")
                st.json(ner_extract(extracted_text))
            else:
                st.error("No text could be extracted.")

        elif ext in ("png", "jpg", "jpeg"):
            pil_img = Image.open(tmp_path).convert("RGB")
            st.image(pil_img, caption="Uploaded image", use_column_width=True)
            with st.spinner("Running OCR..."):
                full_text = ocr_image_pil(pil_img)
            st.subheader("OCR Text")
            st.text(full_text[:3000] + ("..." if len(full_text) > 3000 else ""))
            st.subheader("Rule-based Extraction")
            st.json(remove_null_values(rule_based_financial_entities(full_text)))
            st.subheader("NER Extraction")
            st.json(ner_extract(full_text))

        else:
            st.warning(f"Unsupported file type: {ext}")
    finally:
        os.unlink(tmp_path)

# st.markdown("---")